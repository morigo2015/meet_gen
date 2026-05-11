from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output" / "generated"


COLUMNS = [
    "номер_договору",
    "дата_договору",
    "назва_клієнта",
    "піб_виконавця",
    "адреса_виконавця",
    "податковий_номер",
    "паспорт",
    "телефон",
    "email",
    "банк_виконавця",
    "iban_виконавця",
    "мфо_виконавця",
    "послуги",
    "період_послуг",
    "сума",
    "валюта",
    "акт_номер",
    "дата_акта",
    "рахунок_номер",
    "дата_рахунку",
    "строк_оплати",
    "замовник_назва",
    "замовник_єдрпоу",
    "замовник_адреса",
    "замовник_директор",
    "замовник_iban",
    "замовник_банк",
    "замовник_мфо",
]


ROWS = [
    [
        "SD-2026-001",
        "01.06.2026",
        "Іваненко Олег Петрович",
        "Іваненко Олег Петрович",
        "02094, м. Київ, вул. Попудренка, буд. 18, кв. 42",
        "3012345678",
        "ID-картка № 001234567, орган 8031",
        "+380 67 111 22 33",
        "oleh.ivanenko@example.com",
        "АТ КБ «ПриватБанк»",
        "UA123052990000026005012345678",
        "305299",
        "розробка backend-модуля CRM та API-інтеграцій",
        "червень 2026",
        85000,
        "грн",
        "ACT-2026-001",
        "30.06.2026",
        "INV-2026-001",
        "25.06.2026",
        "5 банківських днів",
        "ТОВ «Софт Лаб Україна»",
        "41234567",
        "01001, м. Київ, вул. Хрещатик, буд. 22, офіс 15",
        "директор Кравченко Анна Сергіївна",
        "UA893052990000026007123456789",
        "АТ КБ «ПриватБанк»",
        "305299",
    ],
    [
        "SD-2026-002",
        "03.06.2026",
        "Петренко Марія Андріївна",
        "Петренко Марія Андріївна",
        "79008, м. Львів, пл. Ринок, буд. 7, кв. 5",
        "2897654321",
        "паспорт МК 123456, виданий Галицьким РВ",
        "+380 50 444 55 66",
        "maria.petrenko@example.com",
        "АТ «Ощадбанк»",
        "UA633226690000026002098765432",
        "322669",
        "розробка frontend-інтерфейсу кабінету користувача",
        "червень 2026",
        64250.50,
        "грн",
        "ACT-2026-002",
        "28.06.2026",
        "INV-2026-002",
        "24.06.2026",
        "7 календарних днів",
        "ТОВ «Софт Лаб Україна»",
        "41234567",
        "01001, м. Київ, вул. Хрещатик, буд. 22, офіс 15",
        "директор Кравченко Анна Сергіївна",
        "UA893052990000026007123456789",
        "АТ КБ «ПриватБанк»",
        "305299",
    ],
    [
        "SD-2026-003",
        "05.06.2026",
        "Шевчук Дмитро Ігорович",
        "Шевчук Дмитро Ігорович",
        "65026, м. Одеса, вул. Дерибасівська, буд. 10, кв. 21",
        "3123456789",
        "ID-картка № 009876543, орган 5112",
        "+380 93 777 88 99",
        "dmytro.shevchuk@example.com",
        "АТ «Універсал Банк»",
        "UA423220010000026009123450001",
        "322001",
        "розробка мобільної інтеграції та технічна підтримка релізу",
        "липень 2026",
        120000,
        "грн",
        "ACT-2026-003",
        "31.07.2026",
        "INV-2026-003",
        "25.07.2026",
        "5 банківських днів",
        "ТОВ «Діджитал Сервіс Груп»",
        "42345678",
        "04070, м. Київ, вул. Верхній Вал, буд. 30",
        "директор Бондар Сергій Миколайович",
        "UA213005280000026000123456789",
        "АТ «ОТП Банк»",
        "300528",
    ],
]


def main() -> None:
    for directory in [
        TEMPLATE_DIR,
        INPUT_DIR,
        OUTPUT_DIR / "contracts",
        OUTPUT_DIR / "acts",
        OUTPUT_DIR / "invoices",
    ]:
        directory.mkdir(parents=True, exist_ok=True)

    create_contract_template(TEMPLATE_DIR / "service_contract.docx")
    create_act_template(TEMPLATE_DIR / "act_acceptance.docx")
    create_invoice_template(TEMPLATE_DIR / "invoice.docx")
    create_workbook(INPUT_DIR / "developer_contractors.xlsx")
    create_readme(BASE_DIR / "README.md")


def create_contract_template(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "ДОГОВІР № {{ номер_договору }}")
    _center(doc, "про надання послуг з розробки програмного забезпечення")
    doc.add_paragraph("м. Київ")
    doc.add_paragraph("Дата укладення: {{ дата_договору }}")
    doc.add_paragraph(
        "{{ замовник_назва }}, код ЄДРПОУ {{ замовник_єдрпоу }}, в особі {{ замовник_директор }}, "
        "що діє на підставі статуту, надалі - Замовник, з однієї сторони, та фізична особа "
        "{{ піб_виконавця }}, реєстраційний номер облікової картки платника податків "
        "{{ податковий_номер }}, надалі - Виконавець, з іншої сторони, уклали цей договір."
    )
    _section(doc, "1. Предмет договору")
    doc.add_paragraph(
        "1.1. Виконавець зобов'язується надати Замовнику послуги з розробки програмного забезпечення: "
        "{{ послуги }}, а Замовник зобов'язується прийняти та оплатити такі послуги."
    )
    doc.add_paragraph("1.2. Період надання послуг: {{ період_послуг }}.")
    _section(doc, "2. Вартість та порядок оплати")
    doc.add_paragraph(
        "2.1. Загальна вартість послуг за цим договором становить {{ сума }} {{ валюта }} без ПДВ."
    )
    doc.add_paragraph(
        "2.2. Оплата здійснюється протягом {{ строк_оплати }} після підписання акта приймання-передачі "
        "та отримання рахунку Виконавця."
    )
    _section(doc, "3. Приймання послуг")
    doc.add_paragraph(
        "3.1. Результат наданих послуг оформлюється актом приймання-передачі. У разі відсутності "
        "письмових зауважень протягом 5 робочих днів послуги вважаються прийнятими."
    )
    _section(doc, "4. Права інтелектуальної власності")
    doc.add_paragraph(
        "4.1. Майнові права інтелектуальної власності на результати, створені Виконавцем у межах "
        "цього договору, переходять до Замовника після повної оплати відповідних послуг."
    )
    _section(doc, "5. Конфіденційність")
    doc.add_paragraph(
        "5.1. Сторони зобов'язуються не розголошувати технічну, комерційну та іншу конфіденційну "
        "інформацію, отриману під час виконання договору."
    )
    _section(doc, "6. Реквізити сторін")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Замовник"
    table.rows[0].cells[1].text = "Виконавець"
    row = table.add_row().cells
    row[0].text = (
        "{{ замовник_назва }}\n"
        "ЄДРПОУ: {{ замовник_єдрпоу }}\n"
        "Адреса: {{ замовник_адреса }}\n"
        "IBAN: {{ замовник_iban }}\n"
        "Банк: {{ замовник_банк }}, МФО {{ замовник_мфо }}"
    )
    row[1].text = (
        "{{ піб_виконавця }}\n"
        "РНОКПП: {{ податковий_номер }}\n"
        "Паспорт: {{ паспорт }}\n"
        "Адреса: {{ адреса_виконавця }}\n"
        "Тел.: {{ телефон }}, email: {{ email }}\n"
        "IBAN: {{ iban_виконавця }}\n"
        "Банк: {{ банк_виконавця }}, МФО {{ мфо_виконавця }}"
    )
    doc.add_paragraph("\nПідпис Замовника: ____________________")
    doc.add_paragraph("Підпис Виконавця: ____________________")
    doc.save(path)


def create_act_template(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "АКТ ПРИЙМАННЯ-ПЕРЕДАЧІ НАДАНИХ ПОСЛУГ № {{ акт_номер }}")
    _center(doc, "до договору № {{ номер_договору }} від {{ дата_договору }}")
    doc.add_paragraph("Дата акта: {{ дата_акта }}")
    doc.add_paragraph(
        "{{ замовник_назва }} в особі {{ замовник_директор }} та фізична особа {{ піб_виконавця }} "
        "склали цей акт про те, що Виконавець надав, а Замовник прийняв такі послуги."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Послуга", "Період", "Вартість", "Примітка"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    row = table.add_row().cells
    row[0].text = "{{ послуги }}"
    row[1].text = "{{ період_послуг }}"
    row[2].text = "{{ сума }} {{ валюта }}"
    row[3].text = "Без ПДВ"
    doc.add_paragraph(
        "Сторони підтверджують, що послуги надані належним чином, претензії щодо обсягу, строків "
        "та якості послуг відсутні."
    )
    doc.add_paragraph("Реквізити Виконавця: {{ піб_виконавця }}, РНОКПП {{ податковий_номер }}, {{ iban_виконавця }}.")
    doc.add_paragraph("\nПідпис Замовника: ____________________")
    doc.add_paragraph("Підпис Виконавця: ____________________")
    doc.save(path)


def create_invoice_template(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "РАХУНОК № {{ рахунок_номер }}")
    doc.add_paragraph("Дата рахунку: {{ дата_рахунку }}")
    doc.add_paragraph("Постачальник: {{ піб_виконавця }}, РНОКПП {{ податковий_номер }}")
    doc.add_paragraph("Адреса постачальника: {{ адреса_виконавця }}")
    doc.add_paragraph("IBAN: {{ iban_виконавця }}")
    doc.add_paragraph("Банк: {{ банк_виконавця }}, МФО {{ мфо_виконавця }}")
    doc.add_paragraph("Платник: {{ замовник_назва }}, ЄДРПОУ {{ замовник_єдрпоу }}")
    doc.add_paragraph("Договір: № {{ номер_договору }} від {{ дата_договору }}")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["№", "Найменування послуг", "Період", "Сума", "ПДВ"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = "{{ послуги }}"
    row[2].text = "{{ період_послуг }}"
    row[3].text = "{{ сума }} {{ валюта }}"
    row[4].text = "Без ПДВ"
    doc.add_paragraph("Разом до сплати: {{ сума }} {{ валюта }} без ПДВ.")
    doc.add_paragraph("Строк оплати: {{ строк_оплати }}.")
    doc.add_paragraph("\nВиконавець: ____________________ / {{ піб_виконавця }} /")
    doc.save(path)


def create_workbook(path: Path) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Дані"
    worksheet.append(COLUMNS)
    for row in ROWS:
        worksheet.append(row)

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in worksheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    widths = [18, 16, 28, 28, 44, 18, 34, 18, 30, 24, 34, 14, 48, 20, 14, 10]
    for index, width in enumerate(widths, start=1):
        worksheet.column_dimensions[get_column_letter(index)].width = width
    for index in range(len(widths) + 1, len(COLUMNS) + 1):
        worksheet.column_dimensions[get_column_letter(index)].width = 24
    workbook.save(path)


def create_readme(path: Path) -> None:
    path.write_text(
        "# Customer Sample: Software Development Services\n\n"
        "This sample set models a Ukrainian software-development service package where a limited "
        "liability company orders development services from an individual contractor.\n\n"
        "Input files:\n"
        "- `input/developer_contractors.xlsx` - contractor and company requisites.\n"
        "- `templates/service_contract.docx` - service contract template.\n"
        "- `templates/act_acceptance.docx` - acceptance act template.\n"
        "- `templates/invoice.docx` - invoice template.\n\n"
        "Recommended generation commands:\n\n"
        "```bash\n"
        "python -m src.main generate --template ./samples/for_customer/templates/service_contract.docx "
        "--data ./samples/for_customer/input/developer_contractors.xlsx "
        "--output ./samples/for_customer/output/generated/contracts "
        "--filename-template \"{номер_договору}_{назва_клієнта}_договір\"\n\n"
        "python -m src.main generate --template ./samples/for_customer/templates/act_acceptance.docx "
        "--data ./samples/for_customer/input/developer_contractors.xlsx "
        "--output ./samples/for_customer/output/generated/acts "
        "--filename-template \"{акт_номер}_{назва_клієнта}_акт\"\n\n"
        "python -m src.main generate --template ./samples/for_customer/templates/invoice.docx "
        "--data ./samples/for_customer/input/developer_contractors.xlsx "
        "--output ./samples/for_customer/output/generated/invoices "
        "--filename-template \"{рахунок_номер}_{назва_клієнта}_рахунок\"\n"
        "```\n\n"
        "The templates are realistic samples for testing document generation and are not legal advice.\n",
        encoding="utf-8",
    )


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _center(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


if __name__ == "__main__":
    main()
