from __future__ import annotations

from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output" / "generated"

COMPANY_A = {
    "замовник_назва": "ТОВ «Софт Лаб Україна»",
    "замовник_єдрпоу": "41234567",
    "замовник_адреса": "01001, м. Київ, вул. Хрещатик, буд. 22, офіс 15",
    "замовник_директор": "Кравченко Анна Сергіївна",
    "замовник_підстава": "Статуту",
    "замовник_iban": "UA893052990000026007123456789",
    "замовник_банк": "АТ КБ «ПриватБанк»",
    "замовник_мфо": "305299",
}

COMPANY_B = {
    "замовник_назва": "ТОВ «Діджитал Сервіс Груп»",
    "замовник_єдрпоу": "42345678",
    "замовник_адреса": "04070, м. Київ, вул. Верхній Вал, буд. 30",
    "замовник_директор": "Бондар Сергій Миколайович",
    "замовник_підстава": "Статуту",
    "замовник_iban": "UA213005280000026000123456789",
    "замовник_банк": "АТ «ОТП Банк»",
    "замовник_мфо": "300528",
}

DEVELOPERS = [
    {
        **COMPANY_A,
        "договір_номер": "FW-DEV-2026-001",
        "договір_дата": "01.06.2026",
        "договір_початок": "01.06.2026",
        "договір_кінець": "31.07.2026",
        "період_робіт": "червень-липень 2026",
        "виконавець_піб": "Іваненко Олег Петрович",
        "виконавець_адреса": "02094, м. Київ, вул. Попудренка, буд. 18, кв. 42",
        "виконавець_податковий_номер": "3012345678",
        "виконавець_паспорт": "ID-картка № 001234567, орган 8031",
        "виконавець_телефон": "+380 67 111 22 33",
        "виконавець_email": "oleh.ivanenko@example.com",
        "виконавець_iban": "UA123052990000026005012345678",
        "виконавець_банк": "АТ КБ «ПриватБанк»",
        "виконавець_мфо": "305299",
        "проект_назва": "CRM Platform API",
        "роль_виконавця": "backend developer",
        "строк_оплати": "5 банківських днів",
    },
    {
        **COMPANY_A,
        "договір_номер": "FW-DEV-2026-002",
        "договір_дата": "01.06.2026",
        "договір_початок": "01.06.2026",
        "договір_кінець": "31.07.2026",
        "період_робіт": "червень-липень 2026",
        "виконавець_піб": "Петренко Марія Андріївна",
        "виконавець_адреса": "79008, м. Львів, пл. Ринок, буд. 7, кв. 5",
        "виконавець_податковий_номер": "2897654321",
        "виконавець_паспорт": "паспорт МК 123456, виданий Галицьким РВ",
        "виконавець_телефон": "+380 50 444 55 66",
        "виконавець_email": "maria.petrenko@example.com",
        "виконавець_iban": "UA633226690000026002098765432",
        "виконавець_банк": "АТ «Ощадбанк»",
        "виконавець_мфо": "322669",
        "проект_назва": "Client Portal UI",
        "роль_виконавця": "frontend developer",
        "строк_оплати": "7 календарних днів",
    },
    {
        **COMPANY_B,
        "договір_номер": "FW-DEV-2026-003",
        "договір_дата": "03.06.2026",
        "договір_початок": "01.06.2026",
        "договір_кінець": "31.07.2026",
        "період_робіт": "червень-липень 2026",
        "виконавець_піб": "Шевчук Дмитро Ігорович",
        "виконавець_адреса": "65026, м. Одеса, вул. Дерибасівська, буд. 10, кв. 21",
        "виконавець_податковий_номер": "3123456789",
        "виконавець_паспорт": "ID-картка № 009876543, орган 5112",
        "виконавець_телефон": "+380 93 777 88 99",
        "виконавець_email": "dmytro.shevchuk@example.com",
        "виконавець_iban": "UA423220010000026009123450001",
        "виконавець_банк": "АТ «Універсал Банк»",
        "виконавець_мфо": "322001",
        "проект_назва": "Mobile Release Automation",
        "роль_виконавця": "mobile integration developer",
        "строк_оплати": "5 банківських днів",
    },
]

MONTHLY_WORK = [
    {
        "developer": 0,
        "month_index": 1,
        "заявка_номер": "REQ-2026-001-06",
        "заявка_дата": "01.06.2026",
        "акт_номер": "ACT-2026-001-06",
        "акт_дата": "30.06.2026",
        "місяць_послуг": "червень 2026",
        "період_початок": "01.06.2026",
        "період_кінець": "30.06.2026",
        "строк_здачі": "30.06.2026",
        "опис_робіт": "розробка API для модулів клієнтів, договорів та подій аудиту",
        "результат_робіт": "реалізовано REST API, міграції бази даних, unit-тести та технічну документацію",
        "критерії_приймання": "код передано у репозиторій Замовника, автоматичні тести проходять, критичних дефектів не виявлено",
        "сума_брутто": Decimal("90000.00"),
    },
    {
        "developer": 0,
        "month_index": 2,
        "заявка_номер": "REQ-2026-001-07",
        "заявка_дата": "01.07.2026",
        "акт_номер": "ACT-2026-001-07",
        "акт_дата": "31.07.2026",
        "місяць_послуг": "липень 2026",
        "період_початок": "01.07.2026",
        "період_кінець": "31.07.2026",
        "строк_здачі": "31.07.2026",
        "опис_робіт": "інтеграція CRM API з білінговою системою та чергою повідомлень",
        "результат_робіт": "підключено інтеграційні endpoints, додано retry-логіку та моніторинг помилок",
        "критерії_приймання": "інтеграційні сценарії погоджено Замовником, smoke-тести проходять",
        "сума_брутто": Decimal("95000.00"),
    },
    {
        "developer": 1,
        "month_index": 1,
        "заявка_номер": "REQ-2026-002-06",
        "заявка_дата": "01.06.2026",
        "акт_номер": "ACT-2026-002-06",
        "акт_дата": "30.06.2026",
        "місяць_послуг": "червень 2026",
        "період_початок": "01.06.2026",
        "період_кінець": "30.06.2026",
        "строк_здачі": "30.06.2026",
        "опис_робіт": "створення інтерфейсу кабінету користувача та форм налаштувань профілю",
        "результат_робіт": "реалізовано React-компоненти, валідацію форм та адаптивну верстку",
        "критерії_приймання": "макети відтворені відповідно до дизайн-системи, дефекти рівня blocker відсутні",
        "сума_брутто": Decimal("78000.50"),
    },
    {
        "developer": 1,
        "month_index": 2,
        "заявка_номер": "REQ-2026-002-07",
        "заявка_дата": "01.07.2026",
        "акт_номер": "ACT-2026-002-07",
        "акт_дата": "31.07.2026",
        "місяць_послуг": "липень 2026",
        "період_початок": "01.07.2026",
        "період_кінець": "31.07.2026",
        "строк_здачі": "31.07.2026",
        "опис_робіт": "доопрацювання UI таблиць, фільтрів, історії змін та повідомлень",
        "результат_робіт": "додано таблиці з пагінацією, фільтри, тости повідомлень та storybook-сторінки",
        "критерії_приймання": "функціональність перевірена на тестовому середовищі, зауваження Замовника усунені",
        "сума_брутто": Decimal("82000.00"),
    },
    {
        "developer": 2,
        "month_index": 1,
        "заявка_номер": "REQ-2026-003-06",
        "заявка_дата": "03.06.2026",
        "акт_номер": "ACT-2026-003-06",
        "акт_дата": "30.06.2026",
        "місяць_послуг": "червень 2026",
        "період_початок": "03.06.2026",
        "період_кінець": "30.06.2026",
        "строк_здачі": "30.06.2026",
        "опис_робіт": "підготовка мобільного релізного pipeline та інтеграція push-повідомлень",
        "результат_робіт": "налаштовано CI pipeline, build signing, push-канали та тестові збірки",
        "критерії_приймання": "тестові збірки встановлюються на пристрої, журнали помилок доступні Замовнику",
        "сума_брутто": Decimal("105000.00"),
    },
    {
        "developer": 2,
        "month_index": 2,
        "заявка_номер": "REQ-2026-003-07",
        "заявка_дата": "01.07.2026",
        "акт_номер": "ACT-2026-003-07",
        "акт_дата": "31.07.2026",
        "місяць_послуг": "липень 2026",
        "період_початок": "01.07.2026",
        "період_кінець": "31.07.2026",
        "строк_здачі": "31.07.2026",
        "опис_робіт": "автоматизація релізних чек-листів, crash reporting та regression checklist",
        "результат_робіт": "додано автоматичні release notes, crash reporting dashboard та контрольні сценарії",
        "критерії_приймання": "релізний пакет сформований, regression checklist погоджений технічним менеджером",
        "сума_брутто": Decimal("112500.75"),
    },
]

CONTRACT_COLUMNS = [
    "договір_номер",
    "договір_дата",
    "договір_початок",
    "договір_кінець",
    "період_робіт",
    "замовник_назва",
    "замовник_єдрпоу",
    "замовник_адреса",
    "замовник_директор",
    "замовник_підстава",
    "замовник_iban",
    "замовник_банк",
    "замовник_мфо",
    "виконавець_піб",
    "виконавець_адреса",
    "виконавець_податковий_номер",
    "виконавець_паспорт",
    "виконавець_телефон",
    "виконавець_email",
    "виконавець_iban",
    "виконавець_банк",
    "виконавець_мфо",
    "проект_назва",
    "роль_виконавця",
    "загальна_сума_брутто",
    "валюта",
    "строк_оплати",
    "ставка_пдфо",
    "ставка_вз",
    "ставка_єсв",
    "податкова_примітка",
]

MONTHLY_COLUMNS = [
    "договір_номер",
    "договір_дата",
    "заявка_номер",
    "заявка_дата",
    "акт_номер",
    "акт_дата",
    "місяць_послуг",
    "період_початок",
    "період_кінець",
    "строк_здачі",
    "замовник_назва",
    "замовник_єдрпоу",
    "замовник_адреса",
    "замовник_директор",
    "замовник_підстава",
    "виконавець_піб",
    "виконавець_адреса",
    "виконавець_податковий_номер",
    "виконавець_паспорт",
    "виконавець_телефон",
    "виконавець_email",
    "виконавець_iban",
    "виконавець_банк",
    "виконавець_мфо",
    "проект_назва",
    "роль_виконавця",
    "опис_робіт",
    "результат_робіт",
    "критерії_приймання",
    "сума_брутто",
    "пдфо_сума",
    "вз_сума",
    "сума_нетто",
    "єсв_сума",
    "валюта",
    "строк_оплати",
    "податкова_примітка",
]


def main() -> None:
    for directory in [
        TEMPLATE_DIR,
        INPUT_DIR,
        OUTPUT_DIR / "contracts",
        OUTPUT_DIR / "work_requests",
        OUTPUT_DIR / "acts",
    ]:
        directory.mkdir(parents=True, exist_ok=True)

    create_framework_contract(TEMPLATE_DIR / "framework_service_contract.docx")
    create_work_request(TEMPLATE_DIR / "monthly_work_request.docx")
    create_acceptance_act(TEMPLATE_DIR / "monthly_acceptance_act.docx")
    create_workbook(INPUT_DIR / "developer_work_2026.xlsx")
    create_readme(BASE_DIR / "README.md")
    create_legal_check(BASE_DIR / "LEGAL_CHECK.md")


def create_framework_contract(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "РАМКОВИЙ ДОГОВІР № {{ договір_номер }}")
    _center(doc, "про надання послуг з розробки програмного забезпечення")
    doc.add_paragraph("Дата укладення: {{ договір_дата }}")
    doc.add_paragraph("Строк дії: з {{ договір_початок }} до {{ договір_кінець }}")
    doc.add_paragraph(
        "{{ замовник_назва }}, код ЄДРПОУ {{ замовник_єдрпоу }}, в особі {{ замовник_директор }}, "
        "що діє на підставі {{ замовник_підстава }}, надалі - Замовник, та фізична особа "
        "{{ виконавець_піб }}, РНОКПП {{ виконавець_податковий_номер }}, надалі - Виконавець, "
        "уклали цей Договір про нижченаведене."
    )

    _section(doc, "1. Предмет і модель співпраці")
    doc.add_paragraph(
        "1.1. Виконавець за завданням Замовника надає послуги з розробки програмного забезпечення "
        "у межах проекту «{{ проект_назва }}» протягом періоду {{ період_робіт }}."
    )
    doc.add_paragraph(
        "1.2. Конкретний обсяг, строк, очікуваний результат, критерії приймання та вартість послуг "
        "визначаються у щомісячній Заявці на виконання робіт, що підписується обома Сторонами."
    )
    doc.add_paragraph(
        "1.3. Цей Договір є цивільно-правовим договором про надання послуг. Він не встановлює трудових "
        "відносин, посади, робочого місця, режиму робочого часу, правил внутрішнього трудового розпорядку "
        "або заробітної плати. Виконавець самостійно організовує надання послуг та не підпорядковується "
        "адміністративним наказам Замовника щодо процесу роботи."
    )

    _section(doc, "2. Заявки на виконання робіт")
    doc.add_paragraph(
        "2.1. Заявка є невід'ємною частиною Договору після підписання Сторонами. У разі розбіжностей "
        "між Заявкою та Договором щодо конкретного місяця пріоритет має Заявка в межах погодженого обсягу."
    )
    doc.add_paragraph(
        "2.2. Орієнтовна роль Виконавця для цілей опису послуг: {{ роль_виконавця }}. Така роль не є "
        "посадою або трудовою функцією."
    )

    _section(doc, "3. Приймання та оплата")
    doc.add_paragraph(
        "3.1. Після завершення місячного періоду Виконавець передає результат послуг, а Сторони "
        "підписують Акт приймання-передачі наданих послуг на підставі відповідної Заявки."
    )
    doc.add_paragraph(
        "3.2. Орієнтовна загальна вартість послуг за строк дії Договору становить {{ загальна_сума_брутто }} "
        "{{ валюта }} до утримання податків і зборів, якщо інше не визначено Заявками."
    )
    doc.add_paragraph(
        "3.3. Оплата здійснюється протягом {{ строк_оплати }} після підписання Акта. {{ податкова_примітка }}"
    )

    _section(doc, "4. Права інтелектуальної власності")
    doc.add_paragraph(
        "4.1. Майнові права інтелектуальної власності на результати, створені Виконавцем за Заявками "
        "до цього Договору, переходять до Замовника у повному складі, на весь строк охорони прав та "
        "на територію всіх країн світу з моменту підписання відповідного Акта та повної оплати послуг."
    )
    doc.add_paragraph(
        "4.2. Передача майнових прав включена до вартості послуг. Особисті немайнові права Виконавця "
        "охороняються відповідно до законодавства України."
    )

    _section(doc, "5. Конфіденційність і дані")
    doc.add_paragraph(
        "5.1. Сторони зобов'язуються зберігати конфіденційність технічної, комерційної, фінансової та "
        "іншої інформації, отриманої при виконанні Договору."
    )
    doc.add_paragraph(
        "5.2. Персональні дані Виконавця обробляються Замовником лише для укладення, виконання, обліку "
        "та звітності за цим Договором."
    )

    _section(doc, "6. Реквізити сторін")
    _party_table(doc)
    doc.add_paragraph("\nПідпис Замовника: ____________________")
    doc.add_paragraph("Підпис Виконавця: ____________________")
    doc.save(path)


def create_work_request(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "ЗАЯВКА НА ВИКОНАННЯ РОБІТ № {{ заявка_номер }}")
    _center(doc, "до рамкового договору № {{ договір_номер }} від {{ договір_дата }}")
    doc.add_paragraph("Дата заявки: {{ заявка_дата }}")
    doc.add_paragraph("Місяць надання послуг: {{ місяць_послуг }}")
    doc.add_paragraph("Період: з {{ період_початок }} до {{ період_кінець }}")
    doc.add_paragraph("Проект: {{ проект_назва }}")
    doc.add_paragraph("Виконавець: {{ виконавець_піб }}, РНОКПП {{ виконавець_податковий_номер }}")
    doc.add_paragraph("Замовник: {{ замовник_назва }}, ЄДРПОУ {{ замовник_єдрпоу }}")

    _section(doc, "1. Завдання на місяць")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Зміст"
    for key, value in [
        ("Опис робіт/послуг", "{{ опис_робіт }}"),
        ("Очікуваний результат", "{{ результат_робіт }}"),
        ("Критерії приймання", "{{ критерії_приймання }}"),
        ("Строк передачі результату", "{{ строк_здачі }}"),
        ("Вартість послуг до утримань", "{{ сума_брутто }} {{ валюта }}"),
    ]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    _section(doc, "2. Умови виконання")
    doc.add_paragraph(
        "2.1. Заявка конкретизує обсяг послуг за Договором і не встановлює посаду, робоче місце, "
        "графік роботи або підпорядкування Виконавця правилам внутрішнього трудового розпорядку."
    )
    doc.add_paragraph(
        "2.2. Результати за цією Заявкою передаються Замовнику для перевірки та приймання за Актом "
        "приймання-передачі наданих послуг."
    )
    doc.add_paragraph("\nПідпис Замовника: ____________________")
    doc.add_paragraph("Підпис Виконавця: ____________________")
    doc.save(path)


def create_acceptance_act(path: Path) -> None:
    doc = Document()
    _set_normal_style(doc)
    _heading(doc, "АКТ ПРИЙМАННЯ-ПЕРЕДАЧІ НАДАНИХ ПОСЛУГ № {{ акт_номер }}")
    _center(doc, "до договору № {{ договір_номер }} від {{ договір_дата }} та заявки № {{ заявка_номер }}")
    doc.add_paragraph("Дата акта: {{ акт_дата }}")
    doc.add_paragraph(
        "{{ замовник_назва }} в особі {{ замовник_директор }} та фізична особа {{ виконавець_піб }} "
        "склали цей Акт про те, що Виконавець надав, а Замовник прийняв послуги за період "
        "{{ період_початок }} - {{ період_кінець }}."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, text in enumerate(["Заявка", "Опис/результат", "Критерії", "Сума"]):
        table.rows[0].cells[index].text = text
    row = table.add_row().cells
    row[0].text = "{{ заявка_номер }}"
    row[1].text = "{{ результат_робіт }}"
    row[2].text = "{{ критерії_приймання }}"
    row[3].text = "{{ сума_брутто }} {{ валюта }}"

    doc.add_paragraph(
        "Сторони підтверджують, що послуги надані в повному обсязі відповідно до Договору та Заявки, "
        "претензії щодо обсягу, строків та якості послуг відсутні."
    )
    doc.add_paragraph(
        "Вартість послуг до утримання податків і зборів: {{ сума_брутто }} {{ валюта }}. "
        "ПДФО: {{ пдфо_сума }} {{ валюта }}; військовий збір: {{ вз_сума }} {{ валюта }}; "
        "сума до виплати Виконавцю: {{ сума_нетто }} {{ валюта }}. ЄСВ нараховується Замовником "
        "додатково у сумі {{ єсв_сума }} {{ валюта }}."
    )
    doc.add_paragraph(
        "Майнові права інтелектуальної власності на результати послуг, зазначені в цьому Акті, "
        "переходять до Замовника відповідно до умов Договору після підписання цього Акта та повної оплати."
    )
    doc.add_paragraph("\nПідпис Замовника: ____________________")
    doc.add_paragraph("Підпис Виконавця: ____________________")
    doc.save(path)


def create_workbook(path: Path) -> None:
    workbook = Workbook()
    contracts_sheet = workbook.active
    contracts_sheet.title = "contracts"
    monthly_sheet = workbook.create_sheet("monthly_work")

    contract_rows = build_contract_rows()
    monthly_rows = build_monthly_rows(contract_rows)
    write_sheet(contracts_sheet, CONTRACT_COLUMNS, contract_rows)
    write_sheet(monthly_sheet, MONTHLY_COLUMNS, monthly_rows)
    workbook.save(path)


def build_contract_rows() -> list[dict[str, object]]:
    monthly_by_developer: dict[int, Decimal] = {0: Decimal("0"), 1: Decimal("0"), 2: Decimal("0")}
    for item in MONTHLY_WORK:
        monthly_by_developer[int(item["developer"])] += item["сума_брутто"]

    rows = []
    for index, developer in enumerate(DEVELOPERS):
        row = {
            **developer,
            "загальна_сума_брутто": monthly_by_developer[index],
            "валюта": "грн",
            "ставка_пдфо": "18%",
            "ставка_вз": "5%",
            "ставка_єсв": "22%",
            "податкова_примітка": (
                "Замовник, як податковий агент при виплаті доходу фізичній особі, утримує ПДФО та "
                "військовий збір із суми винагороди, а ЄСВ нараховує та сплачує додатково відповідно "
                "до чинного законодавства."
            ),
        }
        rows.append(row)
    return rows


def build_monthly_rows(contract_rows: list[dict[str, object]]) -> list[dict[str, object]]:
    rows = []
    for item in MONTHLY_WORK:
        developer_index = int(item["developer"])
        base = contract_rows[developer_index]
        gross = Decimal(item["сума_брутто"])
        pit = money(gross * Decimal("0.18"))
        military = money(gross * Decimal("0.05"))
        net = money(gross - pit - military)
        esv = money(gross * Decimal("0.22"))
        row = {
            **base,
            **{key: value for key, value in item.items() if key not in {"developer", "month_index"}},
            "пдфо_сума": pit,
            "вз_сума": military,
            "сума_нетто": net,
            "єсв_сума": esv,
        }
        rows.append(row)
    return rows


def write_sheet(sheet, columns: list[str], rows: list[dict[str, object]]) -> None:
    sheet.append(columns)
    for row in rows:
        sheet.append([row.get(column, "") for column in columns])
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    for index, column in enumerate(columns, start=1):
        width = min(max(len(column) + 4, 14), 44)
        sheet.column_dimensions[get_column_letter(index)].width = width


def create_readme(path: Path) -> None:
    path.write_text(
        "# Customer Sample 2: Framework Contract + Monthly Requests + Acts\n\n"
        "This sample models a more realistic Ukrainian software-development service workflow:\n\n"
        "- 3 individual developers.\n"
        "- Each developer has one framework service contract for June-July 2026.\n"
        "- Each developer has one monthly work request for June and one for July.\n"
        "- Each monthly act is based on the framework contract and the relevant work request.\n\n"
        "Input files:\n\n"
        "- `input/developer_work_2026.xlsx` with sheets `contracts` and `monthly_work`.\n"
        "- `templates/framework_service_contract.docx`.\n"
        "- `templates/monthly_work_request.docx`.\n"
        "- `templates/monthly_acceptance_act.docx`.\n\n"
        "Generation commands:\n\n"
        "```bash\n"
        "python -m src.main generate --template ./samples/for_customer_2/templates/framework_service_contract.docx "
        "--data ./samples/for_customer_2/input/developer_work_2026.xlsx --sheet contracts "
        "--output ./samples/for_customer_2/output/generated/contracts "
        "--filename-template \"{договір_номер}_{виконавець_піб}_договір\"\n\n"
        "python -m src.main generate --template ./samples/for_customer_2/templates/monthly_work_request.docx "
        "--data ./samples/for_customer_2/input/developer_work_2026.xlsx --sheet monthly_work "
        "--output ./samples/for_customer_2/output/generated/work_requests "
        "--filename-template \"{заявка_номер}_{виконавець_піб}_заявка\"\n\n"
        "python -m src.main generate --template ./samples/for_customer_2/templates/monthly_acceptance_act.docx "
        "--data ./samples/for_customer_2/input/developer_work_2026.xlsx --sheet monthly_work "
        "--output ./samples/for_customer_2/output/generated/acts "
        "--filename-template \"{акт_номер}_{виконавець_піб}_акт\"\n"
        "```\n\n"
        "Expected output: 3 contracts, 6 work requests, and 6 acceptance acts.\n\n"
        "These files are sample templates for document-generation testing and should be reviewed by a "
        "Ukrainian lawyer and accountant before real use.\n",
        encoding="utf-8",
    )


def create_legal_check(path: Path) -> None:
    path.write_text(
        "# Legal Alignment Check\n\n"
        "Checked on 2026-05-12 against current official sources available from Verkhovna Rada and the "
        "State Tax Service. This is a drafting checklist for sample documents, not legal advice.\n\n"
        "Sources checked:\n\n"
        "- Civil Code of Ukraine, current edition shown by Rada as of 2026-02-01: "
        "https://zakon.rada.gov.ua/laws/show/435-15\n"
        "- Labor Code of Ukraine, current edition shown by Rada as of 2026-01-01: "
        "https://zakon.rada.gov.ua/laws/show/322-08\n"
        "- Tax Code of Ukraine, current edition shown by Rada as of 2026-01-01: "
        "https://zakon.rada.gov.ua/laws/show/2755-17\n"
        "- Law of Ukraine on unified social contribution, current edition shown by Rada as of 2026-01-26: "
        "https://zakon.rada.gov.ua/laws/show/2464-17\n"
        "- State Tax Service guidance on civil-law contract remuneration and unified social contribution: "
        "https://kh.tax.gov.ua/media-ark/news-ark/792382.html\n\n"
        "Template alignment notes:\n\n"
        "1. Civil-law service model: the framework contract is structured as a paid service contract, "
        "with monthly requests defining the specific assignment and monthly acts confirming acceptance.\n"
        "2. Monthly request workflow: each request specifies scope, expected result, acceptance criteria, "
        "deadline, and price; the act references both the framework contract and the request.\n"
        "3. Labor-law risk control: the wording avoids job title as employment position, salary, fixed "
        "workplace, fixed working hours, vacation, sick leave, and internal labor rules. It states that "
        "the developer independently organizes service delivery.\n"
        "4. IP rights: the framework contract and acts include assignment of economic IP rights to "
        "ordered software results after act signing and payment, while preserving personal non-property rights.\n"
        "5. Tax/SSC handling: the sample assumes the contractor is an individual who is not an FOP; the "
        "customer acts as tax agent, withholds PIT and military levy from gross remuneration, and accrues "
        "SSC separately. Real cases must confirm contractor status and rates on the payment date.\n"
        "6. Signature flow: every contract, request, and act has signature blocks for both parties.\n\n"
        "Result of template-level check: no structural mismatch found against the above sources for a "
        "sample civil-law service workflow. Real production use still requires fact-specific legal and "
        "tax review.\n",
        encoding="utf-8",
    )


def _party_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Замовник"
    table.rows[0].cells[1].text = "Виконавець"
    row = table.add_row().cells
    row[0].text = (
        "{{ замовник_назва }}\n"
        "ЄДРПОУ: {{ замовник_єдрпоу }}\n"
        "Адреса: {{ замовник_адреса }}\n"
        "IBAN: {{ замовник_iban }}\n"
        "Банк: {{ замовник_банк }}, МФО {{ замовник_мфо }}"
    )
    row[1].text = (
        "{{ виконавець_піб }}\n"
        "РНОКПП: {{ виконавець_податковий_номер }}\n"
        "Паспорт: {{ виконавець_паспорт }}\n"
        "Адреса: {{ виконавець_адреса }}\n"
        "Тел.: {{ виконавець_телефон }}\n"
        "Email: {{ виконавець_email }}\n"
        "IBAN: {{ виконавець_iban }}\n"
        "Банк: {{ виконавець_банк }}, МФО {{ виконавець_мфо }}"
    )


def money(value: Decimal) -> Decimal:
    return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _center(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


if __name__ == "__main__":
    main()
